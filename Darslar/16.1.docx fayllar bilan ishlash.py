from docx import Document

hujjat = Document()

hujjat.add_heading("Salom dunyo", level=3)

hujjat.add_paragraph("Bu matn python kodi orqali yozildi")
hujjat.add_paragraph("Bu matn python kodi orqali yozildi")
hujjat.add_paragraph("Bu matn python kodi orqali yozildi")

# for i in hujjat.paragraphs:
#     print(i.text)

hujjat.save("data/hujjat.docx")
